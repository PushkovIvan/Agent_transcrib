import os
import numpy as np
import torch
import torchaudio
import librosa
import soundfile as sf
from pathlib import Path
from datetime import datetime
from collections import defaultdict
import yaml
from docx import Document
from docx.shared import Inches
from tqdm import tqdm
from sklearn.cluster import AgglomerativeClustering
from pyannote.audio import Pipeline
from pyannote.core import Annotation, Segment
from langchain_gigachat.chat_models import GigaChat
from transformers import Wav2Vec2FeatureExtractor, UniSpeechSatForAudioFrameClassification
from faster_whisper import WhisperModel

# Конфигурация
PATH_TO_CONFIG = "C:\\Users\\Егор\\Documents\\vs_code\\WORK_ML\\Agent_transcrib_new\\models\\diarization_config.yaml"
OUTPUT_RTTM = "result.rttm"


def get_config():
    """Загрузка конфигурационного файла."""
    config_path = os.path.join(os.path.dirname(__file__), '..', 'config.yaml')
    with open(config_path, encoding="utf-8") as f:
        config = yaml.safe_load(f)
    return config

CONFIG = get_config()

giga = GigaChat(
    scope='GIGACHAT_API_CORP',
    credentials=CONFIG["token"]["gigachat"],
    verify_ssl_certs=False,
    model="GigaChat-2-Max",
    temperature=0.7
)

# Кэш моделей
_whisper_model = None
_feature_extractor = None
_embedding_model = None

def load_models():
    """Загрузка моделей для транскрибации и эмбеддингов."""
    global _whisper_model, _feature_extractor, _embedding_model
    if _whisper_model is None:
        _whisper_model = WhisperModel(
            model_size_or_path="Systran/faster-whisper-large-v3",
            device="cpu",
            compute_type="int8",
            cpu_threads=6,
        )
    if _feature_extractor is None:
        _feature_extractor = Wav2Vec2FeatureExtractor.from_pretrained("microsoft/unispeech-sat-base-plus")
    if _embedding_model is None:
        _embedding_model = UniSpeechSatForAudioFrameClassification.from_pretrained("microsoft/unispeech-sat-base-plus")
    return _whisper_model, _feature_extractor, _embedding_model

def generate_detailed_analysis_prompt(transcription_text):
    """Генерация промпта для анализа транскрипции."""
    try:
        prompt_path = os.path.join(os.path.dirname(__file__), '..', 'prompt.txt')
        with open(prompt_path, 'r', encoding='utf-8') as f:
            prompt_template = f.read()
        prompt = prompt_template.format(transcription_text=transcription_text)
        return prompt
    except Exception as e:
        print(f"Ошибка при загрузке промпта: {str(e)}")
        prompt = f"""
Обработай транскрипт рабочего совещания. Сначала попробуй определить участников по контексту (диаризация по наитию), группируя реплики по условным ролям или именам (например: "Иван", "Участник 1", "Менеджер проекта"). Затем выдели: 1) Основные темы/проблемы (кратко суть), 2) Поручения (что, кому/роли, сроки если есть), 3) Ключевые решения, 4) Важные детали (цифры, риски, контакты). Не используй markdown. Если списки - только текстовые обозначения типа "а)" или "-". Группируй информацию по смыслу. Структура ответа: сначала условные говорящие, затем остальные категории. Если диаризация невозможна - пропусти её и выдели только суть.

Транскрипт:
{transcription_text}
"""
        return prompt

def analyze_transcription(text):
    """Анализ транскрипции с помощью GigaChat."""
    try:
        prompt = generate_detailed_analysis_prompt(text)
        analysis_result = giga.invoke(prompt)
        return analysis_result.content
    except Exception as e:
        print(f"Ошибка при анализе текста: {str(e)}")
        return None

def cleanup_files(*files):
    """Удаление временных файлов."""
    for file in files:
        try:
            if file and os.path.exists(file):
                os.remove(file)
        except Exception as e:
            print(f"Не удалось удалить файл {file}: {str(e)}")

def extract_embeddings(audio_path, feature_extractor, model, sr=16000):
    """Извлечение эмбеддингов с помощью UniSpeechSat."""
    try:
        waveform, orig_sr = torchaudio.load(audio_path)
        if orig_sr != sr:
            resampler = torchaudio.transforms.Resample(orig_sr, sr)
            waveform = resampler(waveform)
        inputs = feature_extractor(
            waveform.squeeze().numpy(),
            sampling_rate=sr,
            return_tensors="pt",
            padding=True,
            return_attention_mask=True
        )
        with torch.no_grad():
            outputs = model(**inputs, output_hidden_states=True)
            embeddings = outputs.hidden_states[-1].mean(dim=1).squeeze().numpy()
        return embeddings
    except Exception as e:
        print(f"Ошибка при извлечении эмбеддингов: {str(e)}")
        return None

def segment_audio(audio_path, segment_length=3, overlap=1, sr=16000):
    """Сегментация аудио на фрагменты."""
    try:
        y, sr = librosa.load(audio_path, sr=sr)
        duration = len(y) / sr
        segments = []
        start = 0.0
        while start + segment_length <= duration:
            end = start + segment_length
            segment = y[int(start*sr):int(end*sr)]
            segments.append((start, end, segment))
            start += segment_length - overlap
        return segments
    except Exception as e:
        print(f"Ошибка при сегментации аудио: {str(e)}")
        return None

def cluster_speakers(embeddings, n_speakers=None):
    """Кластеризация говорящих по эмбеддингам."""
    try:
        if n_speakers is None:
            n_speakers = min(5, max(2, len(embeddings) // 3))
        clustering = AgglomerativeClustering(
            n_clusters=n_speakers,
            metric='cosine',
            linkage='average'
        ).fit(embeddings)
        return clustering.labels_
    except Exception as e:
        print(f"Ошибка при кластеризации: {str(e)}")
        return None

def transcribe_with_speakers(whisper_model, audio_path, segments, labels):
    """Транскрипция с определением говорящих (faster-whisper)."""
    try:
        result, info = whisper_model.transcribe(
            audio_path,
            language="ru",
            beam_size=3,
            vad_filter=True,
        )
        tagged_segments = []
        for segment in result["segments"]:
            seg_start = segment["start"]
            seg_end = segment["end"]
            speaker = "Unknown"
            for (start, end, _), label in zip(segments, labels):
                if not (seg_end <= start or seg_start >= end):
                    speaker = f"Speaker {label+1}"
                    break
            tagged_segments.append({
                "start": seg_start,
                "end": seg_end,
                "text": segment["text"],
                "speaker": speaker
            })
        return tagged_segments
    except Exception as e:
        print(f"Ошибка при транскрипции: {str(e)}")
        return None

def format_transcription(tagged_segments):
    """Форматирование результата транскрипции с указанием спикеров."""
    formatted = []
    current_speaker = None
    for seg in tagged_segments:
        if seg["speaker"] != current_speaker:
            formatted.append(f"\n\n**{seg['speaker']}:**\n")
            current_speaker = seg["speaker"]
        formatted.append(seg["text"] + " ")
    return "".join(formatted).strip()

def create_docx_document(transcription_text, analysis_text, filename):
    """Создание Word-документа с транскрипцией и анализом."""
    try:
        doc = Document()
        doc.add_heading('Транскрипция совещания', 0).alignment = 1
        doc.add_paragraph(f'Файл: {os.path.basename(filename)}')
        doc.add_paragraph(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M:%S")}')
        doc.add_paragraph('')
        doc.add_heading('Транскрипция', level=1)
        doc.add_paragraph(transcription_text)
        if analysis_text:
            doc.add_heading('Анализ совещания', level=1)
            doc.add_paragraph(analysis_text)
        recordings_dir = os.path.join(os.path.dirname(__file__), '..', 'recordings')
        base = os.path.splitext(os.path.basename(filename))[0]
        doc_filename = f"{base}_transcription.docx"
        doc_path = os.path.join(recordings_dir, doc_filename)
        doc.save(doc_path)
        print(f"Документ сохранен: {doc_path}")
        return doc_path, None
    except Exception as e:
        print(f"Ошибка при создании документа: {str(e)}")
        return None, str(e)

def process_audio_transcription(audio_path):
    """Основная функция обработки аудио для транскрибации."""
    try:
        print("Загружаем модели...")
        whisper_model, feature_extractor, embedding_model = load_models()
        print("Сегментируем аудио...")
        segments = segment_audio(audio_path)
        if not segments:
            return None, None, "Ошибка при сегментации аудио"
        print("Извлекаем эмбеддинги...")
        embeddings = []
        for start, end, segment in segments:
            segment_path = f"temp_segment_{start:.1f}_{end:.1f}.wav"
            sf.write(segment_path, segment, 48000, 'PCM_24')
            emb = extract_embeddings(segment_path, feature_extractor, embedding_model)
            if emb is not None:
                embeddings.append(emb)
            cleanup_files(segment_path)
        if not embeddings:
            return None, None, "Не удалось извлечь эмбеддинги"
        print("Кластеризуем спикеров...")
        labels = cluster_speakers(np.array(embeddings))
        if labels is None:
            return None, None, "Ошибка при кластеризации спикеров"
        print("Транскрибируем со спикерами...")
        tagged_segments = transcribe_with_speakers(
            whisper_model, audio_path, segments, labels
        )
        if not tagged_segments:
            return None, None, "Ошибка при транскрибации"
        print("Формируем транскрибацию...")
        transcription = format_transcription(tagged_segments)
        print("Вызываем GigaChat для анализа...")
        analysis = analyze_transcription(transcription)
        return transcription, analysis, None
    except Exception as e:
        print(f"Ошибка при обработке аудио: {str(e)}")
        return None, None, str(e)

class change_dir:
    """Контекстный менеджер для временной смены директории."""
    def __init__(self, path):
        self.path = path
        self.original = os.getcwd()
    def __enter__(self):
        os.chdir(self.path)
    def __exit__(self, exc_type, exc_val, exc_tb):
        os.chdir(self.original)

def load_pipeline_from_pretrained(path_to_config: str | Path, device: str = None) -> Pipeline:
    """Загрузка пайплайна pyannote из конфигурации."""
    path_to_config = Path(path_to_config)
    cd_to = path_to_config.parent.parent.resolve()
    with change_dir(cd_to):
        print(f"Загружаем пайплайн из {path_to_config}...")
        pipeline = Pipeline.from_pretrained(path_to_config)
    return pipeline

def diarization_pipeline(audio_file):
    """Диаризация аудио с помощью pyannote."""
    pipeline = load_pipeline_from_pretrained(PATH_TO_CONFIG, device="cuda")
    print(f"Обрабатываем {audio_file}...")
    diarization = pipeline(audio_file)
    with open(OUTPUT_RTTM, "w") as f:
        diarization.write_rttm(f)
    print("\nРезультаты диаризации:")
    audio_segments = ""
    for segment, _, speaker in diarization.itertracks(yield_label=True):
        segment_description = f"[{segment.start:.1f}s - {segment.end:.1f}s] {speaker}"
        print(segment_description)
        audio_segments += segment_description
    return audio_segments

def load_rttm(rttm_path):
    """Загружает RTTM-файл в объект Annotation."""
    annotation = Annotation()
    with open(rttm_path, 'r') as f:
        for line in f:
            parts = line.strip().split()
            if len(parts) < 9:
                continue
            start = float(parts[3])
            duration = float(parts[4])
            end = start + duration
            speaker = parts[7]
            segment = Segment(start, end)
            annotation[segment] = speaker
    return annotation

def transcription_with_diarization_pipeline(audio_path):
    """Транскрипция с диаризацией, используя только faster-whisper."""
    diarization_pipeline(audio_path)
    rttm_path = OUTPUT_RTTM
    # Преобразуем аудио в моно для faster-whisper
    audio, sr = librosa.load(audio_path, sr=16000, mono=True)
    temp_wav = "temp_mono.wav"
    sf.write(temp_wav, audio, sr)
    # Транскрипция через faster-whisper
    whisper_model, _, _ = load_models()
    segments, info = whisper_model.transcribe(temp_wav, language="ru", beam_size=3, vad_filter=True)
    segments = list(segments)  # преобразуем генератор в список
    diarization = load_rttm(rttm_path)
    final_result = []
    for segment in segments:
        if isinstance(segment, dict):
            seg_start = segment['start']
            seg_end = segment['end']
            text = segment['text']
        else:
            seg_start = segment.start
            seg_end = segment.end
            text = segment.text
        seg_duration = seg_end - seg_start
        speaker_stats = defaultdict(float)
        for speech_turn in diarization.itertracks(yield_label=True):
            segment_obj = speech_turn[0]
            speaker = speech_turn[2]
            sp_start = segment_obj.start
            sp_end = segment_obj.end
            overlap_start = max(seg_start, sp_start)
            overlap_end = min(seg_end, sp_end)
            overlap_duration = max(0, overlap_end - overlap_start)
            if overlap_duration > 0:
                coverage = overlap_duration / seg_duration
                speaker_stats[speaker] += coverage
        if speaker_stats:
            best_speaker = max(speaker_stats, key=speaker_stats.get)
            coverage_percent = speaker_stats[best_speaker] * 100
        else:
            best_speaker = "UNKNOWN"
            coverage_percent = 0
        if coverage_percent < 50:
            best_speaker = "UNKNOWN"
        final_result.append({
            "speaker": best_speaker,
            "start": seg_start,
            "end": seg_end,
            "text": text,
            "coverage": f"{coverage_percent:.1f}%"
        })
    # Объединение последовательных реплик одного говорящего
    merged_result = []
    current_speaker = None
    current_text = []
    current_start = 0
    current_end = 0
    for res in final_result:
        if res['speaker'] == current_speaker:
            current_text.append(res['text'])
            current_end = res['end']
        else:
            if current_speaker is not None:
                merged_result.append({
                    "speaker": current_speaker,
                    "start": current_start,
                    "end": current_end,
                    "text": ' '.join(current_text)
                })
            current_speaker = res['speaker']
            current_text = [res['text']]
            current_start = res['start']
            current_end = res['end']
    if current_text:
        merged_result.append({
            "speaker": current_speaker,
            "start": current_start,
            "end": current_end,
            "text": ' '.join(current_text)
        })
    print("\nТранскрипция с диаризацией:")
    transcription = ""
    for i, res in enumerate(merged_result, 1):
        result_description = f"{i}. [{res['speaker']}] ({res['start']:.1f}-{res['end']:.1f}): {res['text']}"
        print(result_description)
        transcription += result_description
    os.remove(temp_wav)
    analysis = analyze_transcription(transcription)
    docx_path, docx_error = create_docx_document(transcription, analysis, audio_path)
    if docx_error:
        return transcription, analysis, f"Ошибка при создании docx: {docx_error}"
    return transcription, analysis, docx_path